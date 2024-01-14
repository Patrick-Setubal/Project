### Importações de Bibliotecas
from docx import Document  # Trabalhar com Word
from python_docx_replace.paragraph import Paragraph  # Objeto Paragrafo
from docx.shared import Inches  # Trabalhar com parametros de imagem do docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Formatar Paragrafos
import re  # Função para localizar strings
import os  # Trabalhar com os arquivos no windows
import pandas as pd
import time


### Funções
# Função editada do "python_docx_replace" p/ ganho de performance
def py_replace(doc, **kwargs: str):
    for p in Paragraph.get_all(doc):
        paragraph = Paragraph(p)
        if "$" in paragraph.p.text:
            find_Keys = re.findall(r"\$\{[\w*\-*\.*]*\}", paragraph.p.text)
            for find_key in find_Keys:
                kwargs[find_key[2:-1]]
                paragraph.replace_key(find_key, str(kwargs[find_key[2:-1]]))


### Variaveis
# Local e texto padrão de onde estão os modelos prontos
Modelos = r"./standard_report.docx"
adress_img = r"./data/img/"
adress_data = r"./data/data.xlsx"


### Automações
def GerarRelatorio(nome, line) -> str:
    global Modelos, adress_img, adress_data

    # Definições iniciais
    try:
        # Pegar Documento Modelo
        doc = Document(Modelos)

        # Definir destino do relatorio
        Adress_Doc_Result = f"./{nome}.docx"

        # Trasnformar DF da line em dict
        Inputs = dict(line)
    except Exception as e:
        f"Erro na Definições iniciais: {str(e)}"

    # Preencher Variaveis do relatorio
    try:
        # Editar todas as variaveis do Documento
        # Utilizar a funcao py_replace que é uma modificação da Bibliooteca "python_docx_replace"
        # Modificada para obter ganho de performance
        py_replace(doc, **Inputs)
    except Exception as e:
        f"Erro Geral do preenchimento do Relatorio: {str(e)}"

    # Inserir imagens no Relatorio
    try:
        # Fluxo para Adicionar Imagem
        for table in doc.tables:
            for row in table.columns:
                for i, cell in enumerate(row.cells):
                    if "$Image" in cell.text:
                        ModelInputs = re.findall(r"\$[\w*\-*\.*]*\$", cell.text)

                        # Percorrer todos os $Imagem..$
                        for ModelInput in ModelInputs:
                            # Procurar nome da imagem e define adress
                            Img = adress_img + str(line[(ModelInput.replace("$", ""))])
                            # Limpar conteudo Escrito
                            cell.text = cell.text.replace(ModelInput, "")

                            # Adicionar imagem e Centralizar
                            try:
                                # Adicionar Imagem
                                cell.add_paragraph().add_run().add_picture(
                                    Img, width=Inches(2.0)
                                )
                                # Centralizar Imagem
                                cell.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Adicionar Nome da Imagem
                                cell.add_paragraph().add_run(
                                    line[(ModelInput.replace("$", ""))][16:]
                                )
                                # Centralizar nome da imagem
                                cell.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                f"Erro ao tentar adicionar a imagem: {str(e)}"
                                cell.text = "Erro"
    except Exception as e:
        f"Erro ao Inserir imagens no Relatorio : {str(e)}"

    # Salvar Documento
    try:
        doc.save(Adress_Doc_Result)
    except Exception as e:
        f"Erro ao Salvar Documento: {str(e)}"


dados = pd.read_excel(adress_data)
print("Start Word Report Generator")
print(f"{len(dados)} reports will be created")

for i, line in dados.iterrows():
    GerarRelatorio(f"Relatorio_{i+1}", line)
    print(f"Relatorio_{i+1} Criado com sucesso")

print("-- Codigo Finalizado -- ")

time.sleep(5)
